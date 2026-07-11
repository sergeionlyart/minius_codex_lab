#!/usr/bin/env python3
from __future__ import annotations

import argparse
import html
import json
import re
import sys
import zipfile
from datetime import UTC, datetime
from html.parser import HTMLParser
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

try:
    from .common import (
        iter_claims,
        load_json,
        safe_html_id,
        safe_word_bookmark,
        sha256_file,
        unit_index,
        write_json,
    )
    from .validate import validate_spec
except ImportError:  # Direct script execution.
    from common import (  # type: ignore[no-redef]
        iter_claims,
        load_json,
        safe_html_id,
        safe_word_bookmark,
        sha256_file,
        unit_index,
        write_json,
    )
    from validate import validate_spec  # type: ignore[no-redef]


CSS = """
:root { color-scheme: light; font-family: Georgia, 'Times New Roman', serif; }
body { max-width: 1040px; margin: 2rem auto; padding: 0 2rem 5rem; line-height: 1.5; }
h1, h2, h3, h4 { line-height: 1.2; }
.meta, .integrity { width: 100%; border-collapse: collapse; margin: 1rem 0 2rem; }
.meta th, .meta td, .integrity th, .integrity td {
  border: 1px solid #777; padding: .45rem; vertical-align: top;
}
.claim { border-left: .3rem solid #555; padding: .65rem .9rem; margin: 1rem 0; }
.claim.assumption { border-style: dashed; }
.claim.low, .claim.not_assessed { background: #f3f3f3; }
.evidence-links { white-space: normal; font-size: .92em; }
.evidence-links a, .backlinks a { margin-right: .45rem; }
.qualification { font-size: .92em; margin-top: .35rem; }
.source { border-top: 2px solid #333; margin-top: 2.5rem; padding-top: 1rem; }
.evidence-unit {
  border: 1px solid #999; padding: .75rem; margin: .85rem 0; scroll-margin-top: 1rem;
}
.evidence-unit:target { outline: 3px solid currentColor; }
.locator, .hash, .status, .source-note { font-size: .86em; }
.context { font-size: .9em; opacity: .82; }
.warning { border: 2px solid #555; padding: .75rem; font-weight: bold; }
code { overflow-wrap: anywhere; }
@media print { body { max-width: none; margin: 0; } a { color: inherit; } }
""".strip()


class _AnchorInspector(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.ids: set[str] = set()
        self.internal_targets: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        values = dict(attrs)
        if values.get("id"):
            self.ids.add(str(values["id"]))
        href = values.get("href")
        if tag == "a" and href and href.startswith("#"):
            self.internal_targets.append(href[1:])


def _slug(value: str) -> str:
    normalized = re.sub(r"[^A-Za-z0-9._-]+", "-", value.strip()).strip("-.")
    return normalized[:100] or "verifiable-document"


def _display_local_path(value: Any) -> str:
    if not isinstance(value, str) or not value:
        return "—"
    path = Path(value)
    if path.is_absolute():
        return "[absolute local path omitted]"
    return value


def _source_metadata_rows(source: dict[str, Any]) -> list[tuple[str, Any]]:
    type_and_jurisdiction = f"{source.get('type', '—')} / {source.get('jurisdiction', '—')}"
    date_version_as_of = (
        f"{source.get('date', '—')} / {source.get('version_date', '—')} / "
        f"{source.get('as_of', '—')}"
    )
    return [
        ("Authority", source.get("authority", "—")),
        ("Identifier", source.get("identifier", "—")),
        ("Type / jurisdiction", type_and_jurisdiction),
        ("Date / version / as-of", date_version_as_of),
        ("Retrieved at", source.get("retrieved_at", "—")),
        ("Official URL", source.get("official_url", "—")),
        ("Local reference", _display_local_path(source.get("local_path"))),
        ("Inclusion mode", source.get("inclusion_mode", "—")),
        ("Source SHA-256", source.get("sha256", "—")),
        ("Limitations", source.get("limitations", "—")),
    ]


def _labels(language: str) -> dict[str, str]:
    if language.lower().startswith("uk"):
        return {
            "as_of": "Станом на",
            "classification": "Класифікація",
            "status": "Статус",
            "verification": "Статус перевірки",
            "evidence": "Доказова прив’язка",
            "counterevidence": "Контрдокази/обмеження",
            "appendices": "Додатки: джерела та точні доказові фрагменти",
            "source": "Джерело",
            "locator": "Локатор",
            "back": "Назад до тез",
            "integrity": "Контроль цілісності",
            "notice": (
                "Автоматична перевірка підтверджує цілісність посилань і хешів, "
                "але не замінює юридичну оцінку достатності та застосовності джерела."
            ),
        }
    return {
        "as_of": "По состоянию на",
        "classification": "Классификация",
        "status": "Статус",
        "verification": "Статус проверки",
        "evidence": "Доказательная привязка",
        "counterevidence": "Контрдоказательства/ограничения",
        "appendices": "Приложения: источники и точные доказательственные фрагменты",
        "source": "Источник",
        "locator": "Локатор",
        "back": "Назад к тезисам",
        "integrity": "Контроль целостности",
        "notice": (
            "Автоматическая проверка подтверждает целостность ссылок и хэшей, "
            "но не заменяет юридическую оценку достаточности и применимости источника."
        ),
    }


def _claim_anchor(claim_id: str) -> str:
    return safe_html_id("claim", claim_id)


def _unit_anchor(unit_id: str) -> str:
    return safe_html_id("evidence", unit_id)


def build_html(spec: dict[str, Any], output_path: Path) -> None:
    document = spec["document"]
    language = str(document.get("language", "uk"))
    labels = _labels(language)
    units = unit_index(spec)
    backlinks: dict[str, list[str]] = {unit_id: [] for unit_id in units}
    for _, claim in iter_claims(spec):
        claim_id = str(claim["claim_id"])
        for ref in claim.get("evidence_refs", []) + claim.get("counterevidence_refs", []):
            if ref in backlinks and claim_id not in backlinks[ref]:
                backlinks[ref].append(claim_id)

    parts = [
        "<!doctype html>",
        f'<html lang="{html.escape(language)}">',
        "<head>",
        '<meta charset="utf-8">',
        '<meta name="viewport" content="width=device-width, initial-scale=1">',
        f"<title>{html.escape(str(document['title']))}</title>",
        f"<style>{CSS}</style>",
        "</head>",
        "<body>",
        f"<h1>{html.escape(str(document['title']))}</h1>",
        '<table class="meta">',
        f"<tr><th>ID</th><td>{html.escape(str(document['document_id']))}</td></tr>",
        f"<tr><th>{labels['as_of']}</th><td>{html.escape(str(document['as_of']))}</td></tr>",
        (
            f"<tr><th>{labels['classification']}</th>"
            f"<td>{html.escape(str(document['classification']))}</td></tr>"
        ),
        (
            f"<tr><th>{labels['status']}</th>"
            f"<td>{html.escape(str(document.get('status', 'draft')))}</td></tr>"
        ),
        "</table>",
        f'<p class="warning">{html.escape(labels["notice"])}</p>',
    ]

    for section in spec["sections"]:
        parts.append(f'<section id="{safe_html_id("section", str(section["section_id"]))}">')
        parts.append(f"<h2>{html.escape(str(section['title']))}</h2>")
        for block in section.get("blocks", []):
            block_type = block.get("type")
            text = html.escape(str(block.get("text", "")))
            if block_type == "heading":
                level = min(max(int(block.get("level", 3)), 2), 4)
                parts.append(f"<h{level}>{text}</h{level}>")
            elif block_type == "paragraph":
                parts.append(f"<p>{text}</p>")
            elif block_type == "claim":
                claim_id = str(block["claim_id"])
                claim_class = " ".join(
                    [
                        "claim",
                        str(block.get("claim_type", "")),
                        str(block.get("confidence", "")),
                    ]
                )
                parts.append(
                    f'<div class="{html.escape(claim_class)}" '
                    f'id="{_claim_anchor(claim_id)}" data-claim-id="{html.escape(claim_id)}">'
                )
                parts.append(
                    f"<p><strong>{html.escape(claim_id)}</strong> "
                    f'<span class="claim-text">{text}</span></p>'
                )
                evidence_refs = block.get("evidence_refs", [])
                counter_refs = block.get("counterevidence_refs", [])
                if evidence_refs:
                    links = []
                    for ref in evidence_refs:
                        source, unit = units[str(ref)]
                        label = f"{source['source_id']}, {unit['locator']}"
                        links.append(
                            f'<a href="#{_unit_anchor(str(ref))}" '
                            f'data-unit-id="{html.escape(str(ref))}">{html.escape(label)}</a>'
                        )
                    parts.append(
                        f'<div class="evidence-links"><strong>{labels["evidence"]}:</strong> '
                        + " ".join(links)
                        + "</div>"
                    )
                if counter_refs:
                    links = []
                    for ref in counter_refs:
                        source, unit = units[str(ref)]
                        label = f"{source['source_id']}, {unit['locator']}"
                        links.append(
                            f'<a href="#{_unit_anchor(str(ref))}" '
                            f'data-unit-id="{html.escape(str(ref))}">{html.escape(label)}</a>'
                        )
                    parts.append(
                        '<div class="evidence-links"><strong>'
                        f"{labels['counterevidence']}:</strong> " + " ".join(links) + "</div>"
                    )
                if block.get("qualification"):
                    parts.append(
                        f'<div class="qualification"><strong>Qualification:</strong> '
                        f"{html.escape(str(block['qualification']))}</div>"
                    )
                parts.append("</div>")
        parts.append("</section>")

    parts.append(f"<h1>{html.escape(labels['appendices'])}</h1>")
    parts.append(f"<p>{html.escape(labels['notice'])}</p>")
    for source in spec["sources"]:
        source_id = str(source["source_id"])
        parts.append(f'<section class="source" id="{safe_html_id("source", source_id)}">')
        parts.append(
            f"<h2>{html.escape(labels['source'])} {html.escape(source_id)} — "
            f"{html.escape(str(source['title']))}</h2>"
        )
        metadata_rows = _source_metadata_rows(source)
        parts.append('<table class="meta">')
        for key, value in metadata_rows:
            parts.append(
                f"<tr><th>{html.escape(str(key))}</th><td>{html.escape(str(value))}</td></tr>"
            )
        parts.append("</table>")
        for unit in source.get("units", []):
            unit_id = str(unit["unit_id"])
            parts.append(
                f'<article class="evidence-unit" id="{_unit_anchor(unit_id)}" '
                f'data-unit-id="{html.escape(unit_id)}">'
            )
            parts.append(
                f'<div class="locator"><strong>{html.escape(unit_id)}</strong> — '
                f"{html.escape(labels['locator'])}: {html.escape(str(unit['locator']))}</div>"
            )
            if unit.get("context_before"):
                parts.append(
                    f'<div class="context">[… {html.escape(str(unit["context_before"]))}]</div>'
                )
            parts.append(f'<p class="evidence-text">{html.escape(str(unit["text"]))}</p>')
            if unit.get("context_after"):
                parts.append(
                    f'<div class="context">[{html.escape(str(unit["context_after"]))} …]</div>'
                )
            parts.append(
                f'<div class="status">{html.escape(labels["verification"])}: '
                f"{html.escape(str(unit['verification_status']))}</div>"
            )
            parts.append(
                '<div class="hash">Unit SHA-256: <code>'
                f"{html.escape(str(unit['sha256']))}</code></div>"
            )
            linked_claims = backlinks.get(unit_id, [])
            if linked_claims:
                links = [
                    f'<a href="#{_claim_anchor(claim_id)}">{html.escape(claim_id)}</a>'
                    for claim_id in linked_claims
                ]
                parts.append(
                    f'<div class="backlinks"><strong>{html.escape(labels["back"])}:</strong> '
                    + " ".join(links)
                    + "</div>"
                )
            parts.append("</article>")
        parts.append("</section>")

    parts.extend(["</body>", "</html>"])
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text("\n".join(parts) + "\n", encoding="utf-8")


def validate_html_links(path: Path) -> list[str]:
    inspector = _AnchorInspector()
    inspector.feed(path.read_text(encoding="utf-8"))
    return sorted(set(inspector.internal_targets) - inspector.ids)


def _add_docx_bookmark(paragraph: Any, name: str, bookmark_id: int) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _add_docx_internal_link(paragraph: Any, anchor: str, text: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _set_docx_font(style: Any, name: str, size_pt: int) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    style.font.name = name
    style.font.size = Pt(size_pt)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def build_docx(spec: dict[str, Any], output_path: Path) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm
    except ImportError as error:
        raise RuntimeError(
            "DOCX build requires python-docx. Install tools/verifiable_document/requirements.txt."
        ) from error

    document_meta = spec["document"]
    language = str(document_meta.get("language", "uk"))
    labels = _labels(language)
    units = unit_index(spec)
    claim_bookmarks = {
        str(claim["claim_id"]): safe_word_bookmark("c", str(claim["claim_id"]))
        for _, claim in iter_claims(spec)
    }
    unit_bookmarks = {unit_id: safe_word_bookmark("e", unit_id) for unit_id in units}
    backlinks: dict[str, list[str]] = {unit_id: [] for unit_id in units}
    for _, claim in iter_claims(spec):
        claim_id = str(claim["claim_id"])
        for ref in claim.get("evidence_refs", []) + claim.get("counterevidence_refs", []):
            if ref in backlinks and claim_id not in backlinks[ref]:
                backlinks[ref].append(claim_id)

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    _set_docx_font(doc.styles["Normal"], "Times New Roman", 14)
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        _set_docx_font(doc.styles[style_name], "Times New Roman", 14)
    normal = doc.styles["Normal"].paragraph_format
    normal.line_spacing = 1.15
    normal.space_before = 0
    normal.space_after = 0
    normal.first_line_indent = Cm(1.25)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(str(document_meta["title"]))

    table = doc.add_table(rows=0, cols=2)
    for key, value in (
        ("ID", document_meta["document_id"]),
        (labels["as_of"], document_meta["as_of"]),
        (labels["classification"], document_meta["classification"]),
        (labels["status"], document_meta.get("status", "draft")),
    ):
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = str(value)
    warning = doc.add_paragraph()
    warning.add_run(labels["notice"]).bold = True

    bookmark_counter = 1
    for section_spec in spec["sections"]:
        doc.add_heading(str(section_spec["title"]), level=1)
        for block in section_spec.get("blocks", []):
            block_type = block.get("type")
            if block_type == "heading":
                level = min(max(int(block.get("level", 2)), 2), 4)
                doc.add_heading(str(block.get("text", "")), level=level)
            elif block_type == "paragraph":
                doc.add_paragraph(str(block.get("text", "")))
            elif block_type == "claim":
                claim_id = str(block["claim_id"])
                paragraph = doc.add_paragraph()
                paragraph.add_run(f"{claim_id}. ").bold = True
                paragraph.add_run(str(block["text"]))
                _add_docx_bookmark(paragraph, claim_bookmarks[claim_id], bookmark_counter)
                bookmark_counter += 1
                evidence_refs = block.get("evidence_refs", [])
                counter_refs = block.get("counterevidence_refs", [])
                if evidence_refs:
                    paragraph.add_run(f" [{labels['evidence']}: ")
                    for index, ref in enumerate(evidence_refs):
                        source, unit = units[str(ref)]
                        if index:
                            paragraph.add_run("; ")
                        _add_docx_internal_link(
                            paragraph,
                            unit_bookmarks[str(ref)],
                            f"{source['source_id']}, {unit['locator']}",
                        )
                    paragraph.add_run("]")
                if counter_refs:
                    paragraph.add_run(f" [{labels['counterevidence']}: ")
                    for index, ref in enumerate(counter_refs):
                        source, unit = units[str(ref)]
                        if index:
                            paragraph.add_run("; ")
                        _add_docx_internal_link(
                            paragraph,
                            unit_bookmarks[str(ref)],
                            f"{source['source_id']}, {unit['locator']}",
                        )
                    paragraph.add_run("]")
                if block.get("qualification"):
                    note = doc.add_paragraph()
                    note.add_run("Qualification: ").bold = True
                    note.add_run(str(block["qualification"]))

    doc.add_heading(labels["appendices"], level=1)
    doc.add_paragraph(labels["notice"])
    for source in spec["sources"]:
        source_id = str(source["source_id"])
        doc.add_heading(f"{labels['source']} {source_id} — {source['title']}", level=2)
        table = doc.add_table(rows=0, cols=2)
        metadata_rows = _source_metadata_rows(source)
        for key, value in metadata_rows:
            cells = table.add_row().cells
            cells[0].text = str(key)
            cells[1].text = str(value)
        for unit in source.get("units", []):
            unit_id = str(unit["unit_id"])
            paragraph = doc.add_paragraph()
            paragraph.add_run(f"{unit_id} — {labels['locator']}: {unit['locator']}").bold = True
            _add_docx_bookmark(paragraph, unit_bookmarks[unit_id], bookmark_counter)
            bookmark_counter += 1
            if unit.get("context_before"):
                doc.add_paragraph(f"[… {unit['context_before']}]")
            doc.add_paragraph(str(unit["text"]))
            if unit.get("context_after"):
                doc.add_paragraph(f"[{unit['context_after']} …]")
            details = doc.add_paragraph()
            details.add_run(f"{labels['verification']}: {unit['verification_status']}; ")
            details.add_run(f"Unit SHA-256: {unit['sha256']}")
            linked_claims = backlinks.get(unit_id, [])
            if linked_claims:
                details.add_run(f"; {labels['back']}: ")
                for index, claim_id in enumerate(linked_claims):
                    if index:
                        details.add_run(", ")
                    _add_docx_internal_link(details, claim_bookmarks[claim_id], claim_id)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def validate_docx_links(path: Path) -> list[str]:
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml")
    root = ElementTree.fromstring(xml)
    bookmarks = {
        element.attrib[f"{{{namespace['w']}}}name"]
        for element in root.findall(".//w:bookmarkStart", namespace)
        if f"{{{namespace['w']}}}name" in element.attrib
    }
    targets = {
        element.attrib[f"{{{namespace['w']}}}anchor"]
        for element in root.findall(".//w:hyperlink", namespace)
        if f"{{{namespace['w']}}}anchor" in element.attrib
    }
    return sorted(targets - bookmarks)


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build a verifiable legal document.")
    parser.add_argument("spec", type=Path, help="Path to the validated JSON spec.")
    parser.add_argument("--out-dir", type=Path, required=True, help="Output directory.")
    parser.add_argument("--html", action="store_true", help="Build self-contained HTML.")
    parser.add_argument("--docx", action="store_true", help="Build DOCX with internal bookmarks.")
    parser.add_argument("--basename", help="Override the output basename.")
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv)
    spec_path = args.spec.resolve()
    try:
        spec = load_json(spec_path)
        validation = validate_spec(spec, spec_path)
    except (OSError, ValueError, json.JSONDecodeError) as error:
        print(f"ERROR: {error}", file=sys.stderr)
        return 2
    if not validation["valid"]:
        print(
            f"ERROR: spec has {validation['summary']['errors']} validation error(s).",
            file=sys.stderr,
        )
        for finding in validation["findings"]:
            if finding["severity"] == "error":
                print(
                    f"  {finding['code']} {finding['path']}: {finding['message']}",
                    file=sys.stderr,
                )
        return 1

    build_html_flag = args.html or not (args.html or args.docx)
    build_docx_flag = args.docx
    output_dir = args.out_dir.resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    basename = _slug(args.basename or str(spec["document"].get("document_id", "document")))
    validation_path = output_dir / f"{basename}.validation.json"
    write_json(validation_path, validation)

    outputs: list[dict[str, Any]] = []
    link_errors: list[str] = []
    try:
        if build_html_flag:
            html_path = output_dir / f"{basename}.html"
            build_html(spec, html_path)
            missing = validate_html_links(html_path)
            link_errors.extend(f"HTML missing target: {target}" for target in missing)
            outputs.append(
                {"format": "html", "path": str(html_path), "sha256": sha256_file(html_path)}
            )
        if build_docx_flag:
            docx_path = output_dir / f"{basename}.docx"
            build_docx(spec, docx_path)
            missing = validate_docx_links(docx_path)
            link_errors.extend(f"DOCX missing bookmark: {target}" for target in missing)
            outputs.append(
                {"format": "docx", "path": str(docx_path), "sha256": sha256_file(docx_path)}
            )
    except (OSError, RuntimeError, KeyError, ValueError, zipfile.BadZipFile) as error:
        print(f"ERROR: {error}", file=sys.stderr)
        return 2

    report = {
        "builder": "minius_codex_lab-verifiable-document/1.0.0-beta.1",
        "built_at_utc": datetime.now(UTC).replace(microsecond=0).isoformat().replace("+00:00", "Z"),
        "spec": str(spec_path),
        "valid": not link_errors,
        "link_errors": link_errors,
        "outputs": outputs,
        "validation_report": str(validation_path),
    }
    report_path = output_dir / f"{basename}.build.json"
    write_json(report_path, report)
    if link_errors:
        for error in link_errors:
            print(f"ERROR: {error}", file=sys.stderr)
        return 1
    for output in outputs:
        print(f"built {output['format']}: {output['path']} sha256={output['sha256']}")
    print(f"validation: {validation_path}")
    print(f"build report: {report_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
