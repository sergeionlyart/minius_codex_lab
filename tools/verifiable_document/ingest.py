#!/usr/bin/env python3
from __future__ import annotations

import argparse
import os
import re
import sys
from collections.abc import Iterable
from dataclasses import dataclass
from datetime import UTC, datetime
from html.parser import HTMLParser
from pathlib import Path

try:
    from .common import normalize_text, sha256_file, sha256_text, write_json
except ImportError:  # Direct script execution.
    from common import (  # type: ignore[no-redef]
        normalize_text,
        sha256_file,
        sha256_text,
        write_json,
    )


@dataclass(frozen=True)
class ExtractedParagraph:
    locator: str
    text: str
    page: int | None = None


class _TextHTMLParser(HTMLParser):
    BLOCK_TAGS = {
        "p",
        "div",
        "li",
        "tr",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
        "blockquote",
        "br",
    }

    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag.lower() in self.BLOCK_TAGS:
            self._parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag.lower() in self.BLOCK_TAGS:
            self._parts.append("\n")

    def handle_data(self, data: str) -> None:
        self._parts.append(data)

    def text(self) -> str:
        return "".join(self._parts)


def _read_text(path: Path) -> str:
    data = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "cp1251"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError(f"Unable to decode text file as UTF-8 or CP1251: {path}")


def _paragraphs_from_text(text: str, locator_prefix: str = "paragraph") -> list[ExtractedParagraph]:
    raw_paragraphs = re.split(r"\n\s*\n|\r\n\s*\r\n", text)
    normalized = [normalize_text(item) for item in raw_paragraphs]
    normalized = [item for item in normalized if item]
    return [
        ExtractedParagraph(locator=f"{locator_prefix} {index}", text=value)
        for index, value in enumerate(normalized, start=1)
    ]


def _extract_txt_or_md(path: Path) -> list[ExtractedParagraph]:
    return _paragraphs_from_text(_read_text(path))


def _extract_html(path: Path) -> list[ExtractedParagraph]:
    parser = _TextHTMLParser()
    parser.feed(_read_text(path))
    return _paragraphs_from_text(parser.text())


def _extract_docx(path: Path) -> list[ExtractedParagraph]:
    try:
        from docx import Document
    except ImportError as error:
        raise RuntimeError(
            "DOCX extraction requires python-docx. Install "
            "tools/verifiable_document/requirements.txt."
        ) from error

    document = Document(path)
    values: list[str] = []
    values.extend(normalize_text(paragraph.text) for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            cell_values = [normalize_text(cell.text) for cell in row.cells]
            if any(cell_values):
                values.append(" | ".join(cell_values))
    values = [value for value in values if value]
    return [
        ExtractedParagraph(locator=f"paragraph/table row {index}", text=value)
        for index, value in enumerate(values, start=1)
    ]


def _extract_pdf(path: Path) -> list[ExtractedParagraph]:
    try:
        from pypdf import PdfReader
    except ImportError as error:
        raise RuntimeError(
            "PDF extraction requires pypdf. Install tools/verifiable_document/requirements.txt."
        ) from error

    reader = PdfReader(path)
    values: list[ExtractedParagraph] = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        page_paragraphs = _paragraphs_from_text(text, locator_prefix="paragraph")
        for paragraph_number, paragraph in enumerate(page_paragraphs, start=1):
            values.append(
                ExtractedParagraph(
                    locator=f"page {page_number}, paragraph {paragraph_number}",
                    text=paragraph.text,
                    page=page_number,
                )
            )
    if not values:
        raise ValueError(
            "No machine-readable text was found in the PDF. Do not invent/OCR silently: "
            "prepare page-image evidence and obtain human verification."
        )
    return values


def extract_paragraphs(path: Path) -> list[ExtractedParagraph]:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".rst", ".csv"}:
        return _extract_txt_or_md(path)
    if suffix in {".html", ".htm"}:
        return _extract_html(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".pdf":
        return _extract_pdf(path)
    raise ValueError(
        f"Unsupported input format {suffix!r}. Use TXT/MD/HTML/DOCX/PDF or extract text explicitly."
    )


def _split_sentences(text: str) -> list[str]:
    boundaries = re.compile(r"(?<=[.!?…])\s+(?=[A-ZА-ЯІЇЄҐ0-9«\"(])")
    values = [normalize_text(value) for value in boundaries.split(text)]
    return [value for value in values if value]


def _expand_units(
    paragraphs: Iterable[ExtractedParagraph], unit_mode: str
) -> list[ExtractedParagraph]:
    if unit_mode == "paragraph":
        return list(paragraphs)
    result: list[ExtractedParagraph] = []
    for paragraph in paragraphs:
        sentences = _split_sentences(paragraph.text)
        for sentence_number, sentence in enumerate(sentences, start=1):
            result.append(
                ExtractedParagraph(
                    locator=f"{paragraph.locator}, sentence {sentence_number}",
                    text=sentence,
                    page=paragraph.page,
                )
            )
    return result


def _relative_reference(input_path: Path, output_path: Path) -> str:
    try:
        return os.path.relpath(input_path.resolve(), output_path.parent.resolve())
    except ValueError:  # Different Windows drives.
        return re.sub(r"[^A-Za-z0-9._-]", "_", input_path.name)


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Create addressable evidence units from a source file."
    )
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--source-id", required=True)
    parser.add_argument("--title", required=True)
    parser.add_argument("--authority", required=True)
    parser.add_argument("--identifier", default="")
    parser.add_argument("--type", default="document")
    parser.add_argument("--jurisdiction", default="Ukraine")
    parser.add_argument("--date", default="")
    parser.add_argument("--version-date", default="")
    parser.add_argument("--as-of", default="")
    parser.add_argument("--effective-from", default="")
    parser.add_argument("--effective-to", default="")
    parser.add_argument("--retrieved-at")
    parser.add_argument("--official-url", default="")
    parser.add_argument(
        "--inclusion-mode",
        choices=("full", "evidence-pack", "page-image-evidence"),
        default="evidence-pack",
    )
    parser.add_argument("--unit-mode", choices=("paragraph", "sentence"), default="paragraph")
    parser.add_argument(
        "--verification-status",
        choices=(
            "human-verified",
            "machine-extracted",
            "needs-human-verification",
            "not-applicable",
        ),
        default="machine-extracted",
    )
    parser.add_argument("--reliability", default="not-assessed")
    parser.add_argument("--limitations", default="")
    parser.add_argument(
        "--local-ref",
        help=(
            "Repository-relative source reference shown in the appendix. Defaults "
            "to a safe relative/basename reference."
        ),
    )
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv)
    source_id = args.source_id.strip()
    if not re.fullmatch(r"[A-Za-z0-9][A-Za-z0-9._-]{0,79}", source_id):
        print(
            "ERROR: --source-id must contain only letters, digits, dot, underscore, or hyphen.",
            file=sys.stderr,
        )
        return 2
    input_path = args.input.resolve()
    output_path = args.output.resolve()
    if not input_path.is_file():
        print(f"ERROR: source file not found: {input_path}", file=sys.stderr)
        return 2
    if args.inclusion_mode == "page-image-evidence" and input_path.suffix.lower() != ".pdf":
        print(
            "ERROR: page-image-evidence ingestion currently requires a PDF with page locators.",
            file=sys.stderr,
        )
        return 2

    try:
        paragraphs = extract_paragraphs(input_path)
        units_raw = _expand_units(paragraphs, args.unit_mode)
    except (OSError, ValueError, RuntimeError) as error:
        print(f"ERROR: {error}", file=sys.stderr)
        return 2
    if not units_raw:
        print("ERROR: no non-empty evidence units were extracted.", file=sys.stderr)
        return 2

    units = []
    for index, extracted in enumerate(units_raw, start=1):
        unit_id = f"{source_id}-U{index:04d}"
        context_before = units_raw[index - 2].text if index > 1 else ""
        context_after = units_raw[index].text if index < len(units_raw) else ""
        unit = {
            "unit_id": unit_id,
            "locator": extracted.locator,
            "text": extracted.text,
            "sha256": sha256_text(extracted.text),
            "verification_status": args.verification_status,
            "context_before": context_before,
            "context_after": context_after,
        }
        if extracted.page is not None:
            unit["page"] = extracted.page
        units.append(unit)

    retrieved_at = args.retrieved_at or (
        datetime.now(UTC).replace(microsecond=0).isoformat().replace("+00:00", "Z")
    )
    source_object = {
        "source_id": source_id,
        "title": args.title,
        "authority": args.authority,
        "identifier": args.identifier,
        "type": args.type,
        "jurisdiction": args.jurisdiction,
        "date": args.date,
        "version_date": args.version_date,
        "as_of": args.as_of,
        "effective_from": args.effective_from,
        "effective_to": args.effective_to or None,
        "retrieved_at": retrieved_at,
        "official_url": args.official_url,
        "local_path": args.local_ref or _relative_reference(input_path, output_path),
        "sha256": sha256_file(input_path),
        "inclusion_mode": args.inclusion_mode,
        "reliability": args.reliability,
        "limitations": args.limitations,
        "units": units,
    }
    write_json(output_path, source_object)
    print(f"wrote {output_path} with {len(units)} unit(s); source sha256={source_object['sha256']}")
    if args.verification_status != "human-verified":
        print("NOTICE: material cited units still require human verification.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
